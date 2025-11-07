"""
Procesor dla dokumentów DOCX (Microsoft Word).
Anonimizuje tekst w paragrafach i tabelach.
"""

import json
import logging
from pathlib import Path
from typing import Dict, Any, Tuple
from datetime import datetime

from docx import Document
from presidio_analyzer import AnalyzerEngine

from app.analyzer import get_supported_entities
from app.anonymizer import anonymize_text, prepare_anonymization_report

logger = logging.getLogger("app.docx_processor")


def process_docx(
    file_path: Path,
    analyzer: AnalyzerEngine,
    config: Dict[str, Any]
) -> Tuple[Path, Dict[str, Any]]:
    """
    Przetwarza dokument DOCX, anonimizując zawartość.

    Args:
        file_path: Ścieżka do pliku DOCX
        analyzer: Skonfigurowany Presidio AnalyzerEngine
        config: Konfiguracja encji

    Returns:
        tuple: (output_path, report) - ścieżka do zanonimizowanego pliku i raport JSON

    Note:
        MVP: Formatowanie może być uproszczone - runs są łączone w jeden tekst.
    """
    logger.info(f"Rozpoczynam przetwarzanie DOCX: {file_path.name}")

    try:
        # Wczytaj dokument
        doc = Document(file_path)
        logger.debug(f"Dokument wczytany. Paragrafów: {len(doc.paragraphs)}, Tabel: {len(doc.tables)}")

        # Zbierz wszystkie wyniki analizy (do raportu)
        all_analyzer_results = []

        # Przetwórz paragrafy
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():  # Pomiń puste paragrafy
                results = _replace_text_in_paragraph(paragraph, analyzer, config)
                all_analyzer_results.extend(results)

        # Przetwórz tabele
        for table_idx, table in enumerate(doc.tables):
            logger.debug(f"Przetwarzanie tabeli {table_idx + 1}/{len(doc.tables)}")
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para in cell.paragraphs:
                        if para.text.strip():
                            results = _replace_text_in_paragraph(para, analyzer, config)
                            all_analyzer_results.extend(results)

        # Przygotuj ścieżkę wyjściową
        output_path = file_path.parent / f"{file_path.stem}.anon{file_path.suffix}"

        # Zapisz zanonimizowany dokument
        doc.save(output_path)
        logger.info(f"Zapisano zanonimizowany dokument: {output_path}")

        # Przygotuj raport
        report = {
            "source_file": str(file_path),
            "output_file": str(output_path),
            "status": "success",
            "timestamp": datetime.now().isoformat(),
            "format": "DOCX",
            "analysis": prepare_anonymization_report(all_analyzer_results, config)
        }

        # Zapisz raport jako osobny plik JSON
        report_path = output_path.with_suffix('.anon.json')
        with open(report_path, 'w', encoding='utf-8') as f:
            json.dump(report, f, ensure_ascii=False, indent=2)

        logger.info(f"Raport zapisany: {report_path}")

        return output_path, report

    except Exception as e:
        logger.exception(f"Błąd podczas przetwarzania {file_path.name}: {e}")
        raise


def _replace_text_in_paragraph(
    paragraph,
    analyzer: AnalyzerEngine,
    config: Dict[str, Any]
):
    """
    Analizuje i anonimizuje tekst w pojedynczym paragrafie.

    Args:
        paragraph: Obiekt paragrafu z python-docx
        analyzer: Presidio AnalyzerEngine
        config: Konfiguracja encji

    Returns:
        list: Wyniki analizy Presidio (RecognizerResult)

    Note:
        MVP: Runs są łączone - formatowanie może się uprościć.
    """
    # Połącz wszystkie runs w jeden tekst
    original_text = paragraph.text

    if not original_text.strip():
        return []

    # Przeanalizuj tekst
    entities = get_supported_entities(config)
    threshold = config.get('threshold', 0.35)

    analyzer_results = analyzer.analyze(
        text=original_text,
        entities=entities,
        language='pl',
        score_threshold=threshold
    )

    if not analyzer_results:
        return []

    # Zastosuj anonimizację
    anonymized_text = anonymize_text(original_text, analyzer_results, config)

    # Wyczyść istniejącą zawartość paragrafu
    paragraph.clear()

    # Wstaw zanonimizowany tekst
    paragraph.add_run(anonymized_text)

    logger.debug(f"Paragraf: {len(analyzer_results)} wykrytych encji")

    return analyzer_results


def _analyze_text(
    analyzer: AnalyzerEngine,
    config: Dict[str, Any],
    text: str
):
    """
    Pomocnicza funkcja do analizy tekstu.

    Args:
        analyzer: Presidio AnalyzerEngine
        config: Konfiguracja
        text: Tekst do analizy

    Returns:
        list: Wyniki analizy Presidio
    """
    if not text.strip():
        return []

    entities = get_supported_entities(config)
    threshold = config.get('threshold', 0.35)

    return analyzer.analyze(
        text=text,
        entities=entities,
        language='pl',
        score_threshold=threshold
    )
